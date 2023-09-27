# -*- codeing = utf-8 -*-
import docx
from docx import Document
import os
import json

imgUrl_list = [
["https://charlie-backend.oss-cn-hongkong.aliyuncs.com/date/%E5%8D%A1%E9%9D%A2/%E4%B8%8E%E4%BD%A0%E5%87%BA%E9%80%83%2B%E4%B8%80%E6%AE%B5.JPG",
"https://charlie-backend.oss-cn-hongkong.aliyuncs.com/date/%E5%8D%A1%E9%9D%A2/%E4%B8%8E%E4%BD%A0%E5%87%BA%E9%80%83%2B%E4%BA%8C%E6%AE%B5.JPG"],
["https://charlie-backend.oss-cn-hongkong.aliyuncs.com/date/%E5%8D%A1%E9%9D%A2/%E4%BB%B2%E5%A4%8F%E9%8E%8F%E9%87%91%2B%E4%B8%80%E6%AE%B5.JPG",
  "https://charlie-backend.oss-cn-hongkong.aliyuncs.com/date/%E5%8D%A1%E9%9D%A2/%E4%BB%B2%E5%A4%8F%E9%8E%8F%E9%87%91%2B%E4%BA%8C%E6%AE%B5.JPG"],
["https://charlie-backend.oss-cn-hongkong.aliyuncs.com/date/%E5%8D%A1%E9%9D%A2/%E5%A4%9C%E5%B9%95%E7%BC%AD%E7%BB%95%2B%E4%B8%80%E6%AE%B5.jpg",
  "https://charlie-backend.oss-cn-hongkong.aliyuncs.com/date/%E5%8D%A1%E9%9D%A2/%E5%A4%9C%E5%B9%95%E7%BC%AD%E7%BB%95%2B%E4%BA%8C%E6%AE%B5.jpg"],
["https://charlie-backend.oss-cn-hongkong.aliyuncs.com/date/%E5%8D%A1%E9%9D%A2/%E5%B0%86%E9%9B%AA%E9%80%BE%E5%86%AC%2B%E4%B8%80%E6%AE%B5.jpg",
  "https://charlie-backend.oss-cn-hongkong.aliyuncs.com/date/%E5%8D%A1%E9%9D%A2/%E5%B0%86%E9%9B%AA%E9%80%BE%E5%86%AC%2B%E4%BA%8C%E6%AE%B5.JPG"],
["https://charlie-backend.oss-cn-hongkong.aliyuncs.com/date/%E5%8D%A1%E9%9D%A2/%E5%BF%83%E5%8C%BF%2B%E4%B8%80%E6%AE%B5.jpg",
"https://charlie-backend.oss-cn-hongkong.aliyuncs.com/date/%E5%8D%A1%E9%9D%A2/%E5%BF%83%E5%8C%BF%2B%E4%BA%8C%E6%AE%B5.jpg"],
    ["https://charlie-backend.oss-cn-hongkong.aliyuncs.com/date/%E5%8D%A1%E9%9D%A2/%E6%97%B6%E4%B8%8E%E7%8E%AB%E7%91%B0%2B%E4%B8%80%E6%AE%B5.JPG",
     "https://charlie-backend.oss-cn-hongkong.aliyuncs.com/date/%E5%8D%A1%E9%9D%A2/%E6%97%B6%E4%B8%8E%E7%8E%AB%E7%91%B0%2B%E4%BA%8C%E6%AE%B5.JPG"],
["https://charlie-backend.oss-cn-hongkong.aliyuncs.com/date/%E5%8D%A1%E9%9D%A2/%E6%98%AF%E6%97%A5%E6%82%A0%E6%82%A0%2B%E4%B8%80%E6%AE%B5.JPG",
 "https://charlie-backend.oss-cn-hongkong.aliyuncs.com/date/%E5%8D%A1%E9%9D%A2/%E6%98%AF%E6%97%A5%E6%82%A0%E6%82%A0%2B%E4%BA%8C%E6%AE%B5.JPG"],
["https://charlie-backend.oss-cn-hongkong.aliyuncs.com/date/%E5%8D%A1%E9%9D%A2/%E6%98%BC%E6%97%A5%E7%96%91%E9%AD%82%2B%E4%B8%80%E6%AE%B5.JPG",
 "https://charlie-backend.oss-cn-hongkong.aliyuncs.com/date/%E5%8D%A1%E9%9D%A2/%E6%98%BC%E6%97%A5%E7%96%91%E9%AD%82%2B%E4%BA%8C%E6%AE%B5.JPG"],
["https://charlie-backend.oss-cn-hongkong.aliyuncs.com/date/%E5%8D%A1%E9%9D%A2/%E7%8B%AE%E5%BF%83%E6%8E%A0%E5%9F%8E%2B%E4%B8%80%E6%AE%B5.jpg",
 "https://charlie-backend.oss-cn-hongkong.aliyuncs.com/date/%E5%8D%A1%E9%9D%A2/%E7%8B%AE%E5%BF%83%E6%8E%A0%E5%9F%8E%2B%E4%BA%8C%E6%AE%B5.jpg"],
["https://charlie-backend.oss-cn-hongkong.aliyuncs.com/date/%E5%8D%A1%E9%9D%A2/%E7%94%9F%E8%80%8C%E6%9C%89%E7%BF%BC%2B%E4%B8%80%E6%AE%B5.JPG",
 "https://charlie-backend.oss-cn-hongkong.aliyuncs.com/date/%E5%8D%A1%E9%9D%A2/%E7%94%9F%E8%80%8C%E6%9C%89%E7%BF%BC%2B%E4%BA%8C%E6%AE%B5.JPG"],
["https://charlie-backend.oss-cn-hongkong.aliyuncs.com/date/%E5%8D%A1%E9%9D%A2/%E7%9C%B8%E4%B8%AD%E7%83%9F%E7%81%AB%2B%E4%B8%80%E6%AE%B5.JPG",
 "https://charlie-backend.oss-cn-hongkong.aliyuncs.com/date/%E5%8D%A1%E9%9D%A2/%E7%9C%B8%E4%B8%AD%E7%83%9F%E7%81%AB%2B%E4%BA%8C%E6%AE%B5.JPG"],
["https://charlie-backend.oss-cn-hongkong.aliyuncs.com/date/%E5%8D%A1%E9%9D%A2/%E7%BC%A0%E7%BB%B5%E6%B8%B8%E6%88%8F%2B%E4%B8%80%E6%AE%B5.JPG",
 "https://charlie-backend.oss-cn-hongkong.aliyuncs.com/date/%E5%8D%A1%E9%9D%A2/%E7%BC%A0%E7%BB%B5%E6%B8%B8%E6%88%8F%2B%E4%BA%8C%E6%AE%B5.JPG"],
["https://charlie-backend.oss-cn-hongkong.aliyuncs.com/date/%E5%8D%A1%E9%9D%A2/%E8%87%B4%E7%A5%9E%E8%B0%95%2B%E4%B8%80%E6%AE%B5.jpg",
 "https://charlie-backend.oss-cn-hongkong.aliyuncs.com/date/%E5%8D%A1%E9%9D%A2/%E8%87%B4%E7%A5%9E%E8%B0%95%2B%E4%BA%8C%E6%AE%B5.jpg"],
["https://charlie-backend.oss-cn-hongkong.aliyuncs.com/date/%E5%8D%A1%E9%9D%A2/%E8%9D%B4%E8%9D%B6%E6%95%88%E5%BA%94%2B%E4%B8%80%E6%AE%B5.JPG",
 "https://charlie-backend.oss-cn-hongkong.aliyuncs.com/date/%E5%8D%A1%E9%9D%A2/%E8%9D%B4%E8%9D%B6%E6%95%88%E5%BA%94%2B%E4%BA%8C%E6%AE%B5.JPG"],
["https://charlie-backend.oss-cn-hongkong.aliyuncs.com/date/%E5%8D%A1%E9%9D%A2/%E9%80%83%E4%BA%A1%E7%AB%A5%E8%AF%9D%2B%E4%B8%80%E6%AE%B5.jpg",
 "https://charlie-backend.oss-cn-hongkong.aliyuncs.com/date/%E5%8D%A1%E9%9D%A2/%E9%80%83%E4%BA%A1%E7%AB%A5%E8%AF%9D%2B%E4%BA%8C%E6%AE%B5.jpg"],
["https://charlie-backend.oss-cn-hongkong.aliyuncs.com/date/%E5%8D%A1%E9%9D%A2/%E9%92%BB%E7%9F%B3%E4%B9%8B%E5%BF%83%2B%E4%B8%80%E6%AE%B5.JPG",
 "https://charlie-backend.oss-cn-hongkong.aliyuncs.com/date/%E5%8D%A1%E9%9D%A2/%E9%92%BB%E7%9F%B3%E4%B9%8B%E5%BF%83%2B%E4%BA%8C%E6%AE%B5.JPG"],
["https://charlie-backend.oss-cn-hongkong.aliyuncs.com/date/%E5%8D%A1%E9%9D%A2/%E9%93%82%E9%87%91%E9%A3%8E%E5%BA%A6%2B%E4%B8%80%E6%AE%B5.JPG",
 "https://charlie-backend.oss-cn-hongkong.aliyuncs.com/date/%E5%8D%A1%E9%9D%A2/%E9%93%82%E9%87%91%E9%A3%8E%E5%BA%A6%2B%E4%BA%8C%E6%AE%B5.JPG"],

]



def extract_content(file_path, json_path):
    # 提取章节号 章节名
    file0 = os.walk(file_path)
    dirpathes = list()
    chap_names = list()  # 章节名
    subchap_names = list()  # 小节名

    for dirpath, dirnames, filenames in file0:  # 从文件夹路径中提取章节信息
        dirpathes.append(dirpath)
        chap_names.append(dirnames)
        subchap_names.append(filenames)

    data = []  # 第一层  每个约会
    date_num = 0

    for cn in chap_names[0]:
        docx_content_path = dirpathes[0] + "\\" + cn + "\\" + "文本"
        # img_path = dirpathes[0] + "\\" + cn + "\\" + "卡面"
        docx_file = os.walk(docx_content_path)
        data.append({
            "name": cn,
            "sub_chaps": []
        })
        sub_chaps = data[date_num]["sub_chaps"]  # 第二层  每小节
        for dirpath, dirnames, filenames in docx_file:  # 从文件夹路径中提取章节信息
            docx_num = 0

            for fn in filenames:
                text_name = fn.split(".")[0].split("-")[1]  # 按照约定格式时 text_name = fn.split("-")[1]..split(".")[0]
                subchap_docx = docx.Document(docx_content_path+"\\"+fn)

                url = subchap_docx.paragraphs[1].text
                sub_chaps.append(
                    {
                        "subchap_name": text_name,
                        "video": url,
                        "para": []
                    }
                )
                docx_content = sub_chaps[docx_num]["para"]  # 第三层 每小节的内容
                speaker = ""
                content = ""
                regular = True
                for paragraph in subchap_docx.paragraphs:
                    line = paragraph.text.strip()
                    if not line:
                        continue
                    if ":" in line:
                        if speaker and content:  # reached start of next dialogue
                            if speaker != "https":
                                if regular:  # 选项前的内容
                                    docx_content.append({
                                        "speaker": speaker,
                                        "content": content,
                                        "tag": tag
                                    })
                                else:  # 选择之后
                                    docx_content.append({
                                        "speaker": speaker,
                                        "content": content,
                                        "tag": tag
                                    })

                        speaker, content = line.split(":")[0], line.split(":")[1]
                        tag = ""
                        if speaker == "查理苏" or speaker == "我" or speaker == "旁白":
                            tag = speaker
                        else:
                            tag = "其他人"
                    elif "分支结束" in line:  # end of a choice
                        if speaker and content:
                            docx_content.append({
                                "speaker": speaker,
                                "content": content,
                                "tag": tag
                            })
                        content = ""
                        docx_content = sub_chaps[docx_num]["para"]
                        regular = True

                    else:  # continuous paragraph
                        content += line + "\n"

                if speaker and content:  # end of doc
                    if regular:
                        docx_content.append({
                            "speaker": speaker,
                            "content": content,
                            "tag": tag
                        })
                    else:
                        docx_content.append({
                            "speaker": speaker,
                            "content": content,
                            "tag": tag
                        })
                docx_num += 1
        date_num += 1
    # print(data)

    # 输出前端规定格式

    dateData = list()
    subContent_num = 0
    dateIndex = 0
    for d in data:
        dateIndex += 1
        dateName = d.get("name")
        dateMenu = list()
        dateUrl = d.get('sub_chaps')[0].get('video')
        for s in d.get('sub_chaps'):
            dateMenu.append(s.get("subchap_name"))
        dateData.append(
            {
                "dateName": dateName,
                "dateImgList": [],
                "dateUrl": dateUrl,
                "dateIndex": str(dateIndex),
                "dateMenu": dateMenu,
                "dateText": []

            }
        )
    text_card = list()  # 按照卡面顺序，小节顺序内容  索引 0-16
    for d in data:
        text_temp_eveycard = list()
        for s in d.get('sub_chaps'):  # 循环每个小节
            text_temp = list()
            for f in range(len(s.get("para"))):
                name = s.get("para")[f].get("speaker")
                content = s.get("para")[f].get("content")
                text_temp.append(   # 取出每个小节的内容，小节按顺序
                    {
                        'name': name,
                        'content': content
                    }
                )
            text_temp_eveycard.append(text_temp)
        text_card.append(text_temp_eveycard)

    card_index = 0
    for dt in dateData:
        text_list = dt.get('dateText')
        menu_num = 0  # 小节的顺序号码
        for menuname in dt.get('dateMenu'):
            menu_num += 1
            text_list.append(
                {
                    menuname: {
                        "dateContent":[]
                    }
                }
            )
        for m in range(len(text_list)):  # 按顺序循环每个小节
            dateContent_list = text_list[m].get(dt.get('dateMenu')[m]).get('dateContent')
            for t in text_card[card_index][m]:
                dateContent_list.append(t)

        # print(card_index)
        img_list = dt.get("dateImgList")
        for cardurl in imgUrl_list[card_index]:
            img_list.append(cardurl)
        card_index += 1






    with open(json_path, "w", encoding="utf-8") as json_file:
        json.dump(dateData, json_file, ensure_ascii=False, indent=4)

    return data


def main():
    path = 'D:\Python\PycharmProjects\CharlieSu\data\文本+图的素材\\约会'
    json_path = "dateeDB.json"
    extract_content(path, json_path)



if __name__ == '__main__':
    main()
