# -*- codeing = utf-8 -*-
import docx
from docx import Document
import os
import json

def extract_content(file_path, json_path):
    # 提取章节号 章节名
    file0 = os.walk(file_path)
    dirpathes = list()
    chap_names = list()  # 章节名
    subchap_names = list()  # 小节名

    for dirpath, dirnames, filenames in file0:  # 从文件夹路径中提取章节信息
        dirpathes.append(dirpath)
        chap_names.append(dirnames)
        subchap_names.append(filenames)

    data = []  # 第一层  每个约会
    date_num = 0
    for cn in chap_names[0]:
        docx_content_path = dirpathes[0] + "\\" + cn + "\\" + "文本"
        docx_file = os.walk(docx_content_path)
        data.append({
            "name": cn,
            "sub_chaps": []
        })
        sub_chaps = data[date_num]["sub_chaps"]  # 第二层  每小节
        for dirpath, dirnames, filenames in docx_file:  # 从文件夹路径中提取章节信息
            docx_num = 0

            for fn in filenames:
                text_name = fn.split(".")[0]  # 按照约定格式时 text_name = fn.split("-)[1]..split(".")[0]
                subchap_docx = docx.Document(docx_content_path+"\\"+fn)

                url = subchap_docx.paragraphs[1].text
                intro = subchap_docx.paragraphs[2].text
                sub_chaps.append(
                    {
                        "subchap_name": text_name,
                        "video": url,
                        "intro": intro,
                        "para": []
                    }
                )
                docx_content = sub_chaps[docx_num]["para"]  # 第三层 每小节的内容
                speaker = ""
                content = ""
                regular = True
                for paragraph in subchap_docx.paragraphs:
                    line = paragraph.text.strip()
                    if not line:
                        continue
                    if ":" in line:
                        if speaker and content:  # reached start of next dialogue
                            if speaker != "https":
                                if regular:  # 选项前的内容
                                    docx_content.append({
                                        "speaker": speaker,
                                        "content": content,
                                        "tag": tag
                                    })
                                else:  # 选择之后
                                    docx_content.append({
                                        "speaker": speaker,
                                        "content": content,
                                        "tag": tag
                                    })

                        speaker, content = line.split(":")[0], line.split(":")[1]
                        tag = ""
                        if speaker == "查理苏" or speaker == "我" or speaker == "旁白":
                            tag = speaker
                        else:
                            tag = "其他人"
                    elif "分支结束" in line:  # end of a choice
                        if speaker and content:
                            docx_content.append({
                                "speaker": speaker,
                                "content": content,
                                "tag": tag
                            })
                        content = ""
                        docx_content = sub_chaps[docx_num]["para"]
                        regular = True

                    else:  # continuous paragraph
                        content += line + "\n"

                if speaker and content:  # end of doc
                    if regular:
                        docx_content.append({
                            "speaker": speaker,
                            "content": content,
                            "tag": tag
                        })
                    else:
                        docx_content.append({
                            "speaker": speaker,
                            "content": content,
                            "tag": tag
                        })
                docx_num += 1
        date_num += 1
    print(data)
    with open(json_path, "w", encoding="utf-8") as json_file:
        json.dump(data, json_file, ensure_ascii=False, indent=4)

    return data


def main():
    # file_path = 'D:\Python\PycharmProjects\CharlieSu\data\倒带轨迹'
    #
    # json_path = "轨迹.json"
    # extract_content(file_path, json_path)
    path = 'D:\Python\PycharmProjects\CharlieSu\data\倒带轨迹'
    file0 = os.walk(path)
    chap_names = list()  # 章节名

    for dirpath, dirnames, filenames in file0:
        chap_names.append(dirnames)

    print(chap_names[0])
    # translator = Translator(from_lang="chinese", to_lang="english")
    for fn in chap_names[0]:
        file_path = 'D:\Python\PycharmProjects\CharlieSu\data\倒带轨迹' + "\\" + fn  # ！！！同倾听
        # englishfn = translator.translate(fn)
        json_path = fn + ".json"
        extract_content(file_path, json_path)

if __name__ == '__main__':
    main()
# -*- codeing = utf-8 -*-
