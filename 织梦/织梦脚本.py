# -*- codeing = utf-8 -*-

import docx
from docx import Document, document
import os
import json


def extract_content(file_path, json_path):
    # 提取章节号 章节名

    file0 = os.walk(file_path)
    dirpathes = list()
    chap_names = list()  # 章节名
    subchap_names = list()  # 小节名

    for dirpath, dirnames, filenames in file0:  # 从文件夹路径中提取章节信息
        dirpathes.append(dirpath)
        chap_names.append(dirnames)
        subchap_names.append(filenames)

    data = list()
    cycles_num = 0
    for i in range(len(dirpathes) - 1):
        path = dirpathes[i + 1]  # 月歌童瑶  月歌童瑶
        for j in range(len(subchap_names[i + 1])):  # [期遇之旅 此刻相逢] [完美主角] [天际漫游 月下相依]
            cycles_num += 1  # 循环次数
            totalNum = 0
            num = 0
            docx_path = path + "\\" + subchap_names[i + 1][j]
            docx_docu = docx.Document(docx_path)

            for paragraph in docx_docu.paragraphs:  # 遍历每一个docx里的每一行
                line = paragraph.text
                if not line:
                    continue
                if "《" in line:  # 提取 docx 标题
                    subchap_name = line.split("《")[1].split("》")[0]
                if "小节" in line:  # 提取 小节号
                    totalNum += 1

            data.append(
                {
                    "chap_name": path.split('•')[1],  # .split('•')[1].split('.')[0]
                    "subchap_name": subchap_name,
                    "totalNum": totalNum,
                    "sections": []
                }
            )
            sections_list = data[cycles_num - 1]["sections"]
            line_num = 0  # 循环小节  前边 小节 后边分支结束
            start_line = []
            end_line = []
            for paragraph in docx_docu.paragraphs:  # 遍历每一个docx里的每一行
                pline = paragraph.text
                if not pline:
                    continue
                if "分支结束" in pline:
                    num += 1
                    sections_list.append(
                        {
                            'num': num,
                            'cardText': []
                        }
                    )

                line_num += 1  # 此行所在行号
                if "小节" in pline:  # section开始
                    start_line.append(line_num + 1)
                elif "分支结束" in pline:  # section结束
                    end_line.append(line_num - 1)

            for k in range(num):  # 4466  k 是当前所在的小节数 num 在这里是总小节数
                cardText_list = sections_list[k]["cardText"]
                # 读取每小节文本
                contentType = ""
                content = ""

                for n in range(start_line[k], end_line[k] + 1):
                    sline = docx_docu.paragraphs[n].text.strip()
                    if ":" in sline:
                        if contentType and content:
                            cardText_list.append(
                                {
                                    "type": contentType,
                                    "content": content
                                }
                            )

                        contentType, content = sline.split(":")[0], sline.split(":")[1]
                    else:  # continuous paragraph
                        content += sline + "\n"

    print(data)
    # 输出前端规定格式
    contentText = list()
    for d in data:
        if d.get("chap_name") == "月歌童瑶":
            if d.get("subchap_name") == "期遇之旅":
                for s in d.get("sections"):
                    currentPage = s.get("num")
                    contentText.append(
                        {
                            'cardindex': 0,
                            'sessionIndex': 0,
                            'currentPage': currentPage,
                            'subContent': []
                        }
                )
                    subContent_list = contentText[currentPage-1]['subContent']
                    for ct in s.get("cardText"):
                        subContent_list.append(
                            {
                                "type": ct.get("type"),
                                "content": ct.get("content")
                            }
                        )
            else:
                for s in d.get("sections"):
                    currentPage = s.get("num")
                    contentText.append(
                        {
                            'cardindex': 0,
                            'sessionIndex': 1,
                            'currentPage': currentPage,
                            'subContent': []
                        }
                    )
                    subContent_list = contentText[currentPage + 4 - 1]['subContent']
                    for ct in s.get("cardText"):
                        subContent_list.append(
                            {
                                "type": ct.get("type"),
                                "content": ct.get("content")
                            }
                        )
        elif d.get("chap_name") == "月升怪谈":  # 月升
            if d.get("subchap_name") == "天际漫游":
                for s in d.get("sections"):
                    currentPage = s.get("num")
                    contentText.append(
                        {
                            'cardindex': 1,
                            'sessionIndex': 0,
                            'currentPage': currentPage,
                            'subContent': []
                        }
                    )
                    subContent_list = contentText[currentPage + 14 - 1]['subContent']  # 4+4+6
                    for ct in s.get("cardText"):
                        subContent_list.append(
                            {
                                "type": ct.get("type"),
                                "content": ct.get("content")
                            }
                        )

            else:
                for s in d.get("sections"):
                    currentPage = s.get("num")
                    contentText.append(
                        {
                            'cardindex': 1,
                            'sessionIndex': 1,
                            'currentPage': currentPage,
                            'subContent': []
                        }
                    )
                    subContent_list = contentText[currentPage + 20 - 1]['subContent']  # 4+4+6+6
                    for ct in s.get("cardText"):
                        subContent_list.append(
                            {
                                "type": ct.get("type"),
                                "content": ct.get("content")
                            }
                        )
        elif d.get("chap_name") == "假面舞会":  #
            if d.get("subchap_name") == "完美主角":
                for s in d.get("sections"):
                    currentPage = s.get("num")
                    contentText.append(
                        {
                            'cardindex': 2,
                            'sessionIndex': 0,
                            'currentPage': currentPage,
                            'subContent': []
                        }
                    )
                    subContent_list = contentText[currentPage + 8 - 1]['subContent']  # 4+4
                    for ct in s.get("cardText"):
                        subContent_list.append(
                            {
                                "type": ct.get("type"),
                                "content": ct.get("content")
                            }
                        )



    print(contentText)
    with open(json_path, "w", encoding="utf-8") as json_file:
        json.dump(contentText, json_file, ensure_ascii=False, indent=4)

    return data


def main():
    # 邂逅数据文件夹所在位置
    file_path = 'D:\Python\PycharmProjects\CharlieSu\data\织梦'  # 请将此路径替换为您的docx文件路径

    json_path = "dream_weavingDB_update.json"
    extract_content(file_path, json_path)


if __name__ == '__main__':
    main()
