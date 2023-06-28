import csv
from docx import Document
import os
import json
import re

global subchap_data_list 
subchap_data_list  = []

global chap_data_list
chap_data_list = []

# 提炼小节文本
def extract_content(chap_num, sub_type, docx_path, json_path):

    document = Document(docx_path)
    #json 
    #提取file name
    file_name = os.path.basename(docx_path)
    file_name_without_extension = os.path.splitext(file_name)[0]

    if sub_type != '普通':
        file_name_without_extension = chap_num + '-' + file_name_without_extension

    # with open(json_path, 'r', encoding="utf-8") as file:
    #     json_data = json.load(file)

    data = {
        "subchap_name": file_name_without_extension,
        "subchap_type": sub_type,
        "para": []
    } 

    choice = {
        "para_type": "choice",
        "xuanxiang": []
    }

    speaker = ""
    content = []
    current_list = data["para"]
    regular = True 

    for paragraph in document.paragraphs:
        line = paragraph.text.strip()

        if not line:
            continue

        if ":" in line:
            if speaker and content: #reached start of next dialogue 
                if regular:
                    current_list.append({
                        "para_type": "normal",
                        "speaker": speaker,
                        "content": content,
                        "tag": tag
                    })    
                else:
                    current_list.append({
                        "speaker": speaker,
                        "content": content,
                        "tag": tag
                    })   
            
            speaker,choice_name = line.split(":")[0], line.split(":")[1]
            content = []
            tag = ""

            if speaker == "选项光":
                light = {
                    "para_type": "light",
                    "choice_name": choice_name,
                    "choice_para": []
                }
                # data["para"].append(light)
                choice["xuanxiang"].append(light)
                current_list = light["choice_para"]
                regular = False
                content = []    
               
            elif speaker == "选项夜":
                night = {
                    "para_type": "night",
                    "choice_name": choice_name,
                    "choice_para": []
                }
                # data["para"].append(night)
                choice["xuanxiang"].append(night)
                current_list = night["choice_para"]
                regular = False
                content = []   
            #处理 tag 
          
            elif speaker == "查理苏":
                tag = "charlie"
            elif speaker == "我":
                tag = "me"
            elif speaker == "旁白":
                tag = "pb"
            else:
                tag = "others"

               
        elif "分支结束" in line: #end of a choice
            if speaker and content:
                current_list.append({
                    "speaker": speaker,
                    "content": content,
                    "tag": tag
                })
            data["para"].append(choice)
            choice = {
                "para_type": "choice",
                "xuanxiang": []
            }
            content = [] 
            current_list = data["para"]
            regular = True 

        else: #continuous paragraph 
            content.append(line)
            
    if speaker and content: #end of doc
        if regular:
            current_list.append({
                "para_type": "normal",
                "speaker": speaker,
                "content": content,
                "tag": tag
            })    
        else:
            current_list.append({
                "speaker": speaker,
                "content": content,
                "tag": tag
            })    

    # json_data = json.dumps(data, ensure_ascii=False).encode('utf-8')
    subchap_data_list.append(data)

    with open(json_path, "w", encoding="utf-8") as json_file:
        json.dump(subchap_data_list, json_file, ensure_ascii=False, indent=4)

    return file_name_without_extension

    # print(json_data.decode('utf-8'))

    # return data

# 整合大章节信息
def extract_chap(docx_path, item, data):
    # print(docx_path)
    document = Document(docx_path)

    if item == '简介.docx': #简介
        
        for paragraph in document.paragraphs:
            line = paragraph.text.strip()

            if not line:
                continue

            if ":" in line:
                speaker, content = line.split(":")[0], line.split(":")[1]
                if speaker == "章节名":
                    data["name"] = content
                elif speaker == "简介":
                    data["intro"] = content
                else:
                    data["video"] = line
        

    else: #幕后
        sub_behind = { 
            "behind_name" : "",
            "content" : []
        }
        file_name_without_extension = os.path.splitext(item)[0]
        # print(file_name_without_extension)
        sub_behind["behind_name"] = file_name_without_extension

        for paragraph in document.paragraphs:
            line = paragraph.text.strip()
            sub_behind["content"].append(line)
        
        data["behind"].append(sub_behind)
        
        # with open(json_path, "w", encoding="utf-8") as json_file:
        #     json.dump(chap_data_list, json_file, ensure_ascii=False, indent=4)
        
    return data

def sort_by_integer(filename):
    # 使用正则表达式提取文件名中的整数部分
    match = re.match(r'(\d+)-', filename)
    if match:
        number = int(match.group(1))
        return number
    return 0  # 如果文件名不符合格式要求，则返回 0 进行排序


def main():
    os.chdir('./主线文本') #mark data as root dir

    chapters = sorted(os.listdir(),key=sort_by_integer) #find all subdirs / chapters & sort

    print(chapters)
    for chaps in chapters: 
        chap_path = './' + chaps
        chap_num, chap_name = chaps.split("-")
        data = {
                "name" : "", 
                "chap_num" : int(chap_num), 
                "chap_name" : chap_name,
                "intro" : "", 
                "image" : "", 
                "video" : "", 
                "behind" : [], 
                # "subchap" : [] 
        }
        print(chaps)
        subchap_nums = [] # 储存当前所有小节的编号

        for item in os.listdir(chap_path): #每章节里面的所有文档/文件夹
            sub_path = os.path.join(chap_path, item) 

            if os.path.isdir(sub_path): #dirs
                json_path = "../subchaps.json" 
                subchaps = os.listdir(sub_path) #所在文件夹的所有小节
                
                
                if item == "光夜选择" or item == "普通主线":
                    for sub in subchaps: #每小节文档
                        curr_path = os.path.join(sub_path, sub) 
                        subchap_nums.append(extract_content(chap_num, '普通', curr_path, json_path))


                elif item == "光夜结局":
                    for sub in subchaps: #每小节文档
                        curr_path = os.path.join(sub_path, sub) 
                        subchap_nums.append(extract_content(chap_num, sub.split(".")[0], curr_path, json_path))

            elif os.path.isfile(sub_path): #docs 
                json_path = "../chaps.json"
                data = extract_chap(sub_path, item, data)
                # 把subchap list加入data
                data["subchap"] = subchap_nums

        # 生成大章节json文件
        chap_data_list.append(data)
        with open(json_path, "w", encoding="utf-8") as json_file:
            json.dump(chap_data_list, json_file, ensure_ascii=False, indent=4)

                    

    # filename = './主线文本/7-13'  # 请将此路径替换为您的docx文件路径
    # docx_path = filename + '.docx'
    
    # json_path = "test" + ".json"
    # extract_content(docx_path, json_path)
    
    

if __name__ == '__main__':
    main()
