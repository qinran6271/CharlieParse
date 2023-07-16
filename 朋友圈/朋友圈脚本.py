import csv
from docx import Document
import os
import json
import re
import uuid

global details_data_list 
details_data_list  = []

global overview_data_list
overview_data_list = []

# 提炼小节文本
def extract_details(indexCode, docx_path, json_path):

    document = Document(docx_path)

    #json 
    #提取file name, used for pics 
    file_name = os.path.basename(docx_path)
    file_name_without_extension = os.path.splitext(file_name)[0]
    # print(file_name,file_name_without_extension)

    # with open(json_path, 'r', encoding="utf-8") as file:
    #     json_data = json.load(file)

    data = {
        'indexCode':indexCode, 
        'postPerson':"", 
        'postText':"", 
        'postImg' :"", 
        'hasImg': False, 
        'commentPerson': "",
        'commentChoices': [],
        'otherComment':[ ]
    } 

    choice = {
        'index': 0,
        'choiceContent':"",
        'reply':{
            'person':"",
            'content':""
            }
    }

    speaker = ""
    content = ""
    # current_list = data["para"]
    isPost = True 
    paragraphs = [] 
    
    for paragraph in document.paragraphs:
        text_parts = paragraph.text.split('\n')
        # print(text_parts)
        for text_part in text_parts:
                # 忽略空文本部分
            paragraphs.append(text_part.strip())

    for line in paragraphs:
        
    # for paragraph in document.paragraphs:
    #     line = paragraph.text.strip()
        
            # print(line)

        if not line:
            continue

        if ":" in line:
            if speaker and content != "": #reached start of next dialogue 
                # print(speaker,content)
                if isPost: #朋友圈内容
                    data["postPerson"] = speaker
                    data["postText"] = content
                    isPost = False
                else: #评论内容
                    if speaker == "我": 
                        data["commentPerson"] = speaker
                        choice['choiceContent'] = content
                    elif "回复" in speaker: #查理苏回复我
                        choice['reply']['person'] = speaker.split("回复")[0]
                        choice['reply']['content'] = content
                        data['commentChoices'].append(choice)
                    else: #其他人回复
                        data['otherComment'].append({
                            'name': speaker,
                            'content' : content
                        })
                        

            speaker,choice_name = line.split(":")[0], line.split(":")[1]
            content = ""

            if speaker == "朋友圈内容":
                isPost = True

            elif speaker == "照片":
                if choice_name == "是":
                    folder_path = '../照片'  # 替换为实际文件夹的路径

                    for file in os.listdir(folder_path):
                        # 提取文件名（不包含路径和后缀）
                        base_name = os.path.splitext(os.path.basename(file))[0]

                        if base_name == file_name_without_extension:
                            data["postImg"] = "https://charlie-backend.oss-cn-hongkong.aliyuncs.com/moments/" + file
                            break

                    data["hasImg"] = True
                    
                else:
                    data["hasImg"] = False
            
            # elif speaker == "评论内容":
            #     isPost = False

            elif speaker == "choice":
                choice = {
                    'index': int(choice_name),
                    'choiceContent':"",
                    'reply':{
                        'person':"",
                        'content':""
                    }
                }

        else: #continuous paragraph 
            content = line
            # print(speaker,content)
            
    if speaker and content != "": #end of doc
        if speaker == "我": 
            data["commentPerson"] = speaker
            choice['choiceContent'] = content
        elif "回复" in speaker: #查理苏回复我
            choice['reply']['person'] = speaker.split("回复")[0]
            choice['reply']['content'] = content
            data['commentChoices'].append(choice)
        else: #其他人回复
            data['otherComment'].append({
                'name': speaker,
                'content' : content
            })
            

    # json_data = json.dumps(data, ensure_ascii=False).encode('utf-8')
    # print(data)
    details_data_list.append(data)

    with open(json_path, "w", encoding="utf-8") as json_file:
        json.dump(details_data_list, json_file, ensure_ascii=False, indent=4)

    return file_name_without_extension

    # print(json_data.decode('utf-8'))

    # return data


def sort_by_integer(filename):
    # 使用正则表达式提取文件名中的整数部分
    # match = re.match(r'(\d+)-', filename)
    # if match:
    #     number = int(match.group(1))
    #     return number
    # return 0  # 如果文件名不符合格式要求，则返回 0 进行排序
    match = re.search(r'(\d+(\.\d+)?)', filename)
    if match:
        number = float(match.group(1))
        return number
    return 0.0  # 如果文件名不包含数字，则返回 0.0 进行排序

def parse_filename(file_name):
    parts = file_name.split('-')
    if len(parts) == 3:
        category = parts[1]
        content = parts[2]
    elif len(parts) == 2:
        category = parts[0]
        content = parts[1]
    else:
        category = None
        content = None
    return category,content

def main():
    os.chdir('./朋友圈/朋友圈文本1')
    types = os.listdir()
    overview_json_path = "../朋友圈overview.json"
    for type_name in types: 
        type_path = './' + type_name

        # 替换成英语名称 类型: lingXi（灵犀）、xieHou（邂逅）、activities（活动）、truthorDare（真话冒险）、teaParty（茶歇）、mainStory（主线）
        if type_name == "灵犀":
            type_name = "lingXi"
        elif type_name == "邂逅":
            type_name = "xieHou"
        elif type_name == "活动":
            type_name = "activities"
        elif type_name == "真心话大冒险":
            type_name = "truthorDare"
        elif type_name == "茶歇小憩":
            type_name = "teaParty"   
        else:
            type_name = "mainStory"  

        overview = {
                'type': type_name, #类型: lingXi（灵犀）、xieHou（邂逅）、activities（活动）、truthorDare（真话冒险）、teaParty（茶歇）、mainStory（主线）
                
                'data':[],
            } 

        # type_list = os.listdir(type_path)

        # 文件夹需要排序
        type_list = sorted(os.listdir(type_path),key=sort_by_integer)

        for class_name  in type_list: #每章节里面的所有文档/文件夹

            
            print(class_name)
            overview_data = {
                    'className':class_name, #大类名称 
                    'items':[]
                }
            
            if '灵犀' in class_name:
                overview_data["className"] = class_name[2:]
            
            sub_path = os.path.join(type_path, class_name) # sub_path 是每个type里面文件的路径

            sub_list = sorted(os.listdir(sub_path),key=sort_by_integer)
            for item in sub_list:
                print(item)
                
                # doc_path
                cur_path = os.path.join(sub_path, item)
                # 生成indexcode
                unique_code = str(uuid.uuid4())
                indexCode = "wm"+ unique_code
                #json_path
                json_path = "../朋友圈details.json"

                #提取file name
                file_name = os.path.basename(cur_path)
                file_name_without_extension = os.path.splitext(file_name)[0]

                item = {
                    'name':parse_filename(file_name_without_extension)[0], #获取途径
                    'content':parse_filename(file_name_without_extension)[1], #朋友圈简介内容
                    'indexCode':indexCode #朋友圈索引
                }
                overview_data["items"].append(item)
                extract_details(indexCode, cur_path, json_path)
            overview["data"].append(overview_data)


        overview_data_list.append(overview)
        with open(overview_json_path, "w", encoding="utf-8") as json_file:
            json.dump(overview_data_list, json_file, ensure_ascii=False, indent=4)

    # doc_path = './10-第十章-装修新房.docx'  # 请将此路径替换为您的docx文件路径
    
    # json_path = "10-第十章-装修新房.json"

    # indexCode = "wm0001"
    # extract_details(indexCode, doc_path, json_path)
    
    

if __name__ == '__main__':
    main()
